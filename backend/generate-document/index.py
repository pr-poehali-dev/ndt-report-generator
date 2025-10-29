import json
import base64
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from typing import Dict, Any

def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    '''
    Business: Generate NDT conclusion DOCX document from form data
    Args: event - dict with httpMethod, body containing conclusion data
          context - object with attributes: request_id, function_name
    Returns: HTTP response with base64-encoded DOCX file
    '''
    method: str = event.get('httpMethod', 'GET')
    
    if method == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'POST, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type, X-User-Id',
                'Access-Control-Max-Age': '86400'
            },
            'body': ''
        }
    
    if method != 'POST':
        return {
            'statusCode': 405,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*'
            },
            'body': json.dumps({'error': 'Method not allowed'})
        }
    
    body_data = json.loads(event.get('body', '{}'))
    
    joint_number = body_data.get('jointNumber', '')
    inspector_name = body_data.get('inspectorName', '')
    date = body_data.get('date', '')
    control_type = body_data.get('controlType', '')
    results = body_data.get('results', '')
    notes = body_data.get('notes', '')
    
    control_type_names = {
        'ut': 'Ультразвуковой контроль (УЗК)',
        'rt': 'Радиографический контроль (РК)',
        'mt': 'Магнитопорошковый контроль (МПК)',
        'pt': 'Капиллярный контроль (ПВК)',
        'vt': 'Визуальный контроль (ВИК)'
    }
    control_type_full = control_type_names.get(control_type, control_type)
    
    doc = Document()
    
    title = doc.add_heading('ЗАКЛЮЧЕНИЕ', level=1)
    title.alignment = 1
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    subtitle = doc.add_heading('по результатам неразрушающего контроля', level=2)
    subtitle.alignment = 1
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    p1 = doc.add_paragraph()
    p1.add_run('Номер стыка: ').bold = True
    p1.add_run(joint_number)
    
    p2 = doc.add_paragraph()
    p2.add_run('Дата проведения контроля: ').bold = True
    p2.add_run(date)
    
    p3 = doc.add_paragraph()
    p3.add_run('Тип контроля: ').bold = True
    p3.add_run(control_type_full)
    
    p4 = doc.add_paragraph()
    p4.add_run('Дефектоскопист: ').bold = True
    p4.add_run(inspector_name)
    
    doc.add_paragraph()
    
    results_heading = doc.add_heading('Результаты контроля:', level=2)
    results_heading_run = results_heading.runs[0]
    results_heading_run.font.size = Pt(12)
    
    results_para = doc.add_paragraph(results)
    
    if notes:
        doc.add_paragraph()
        notes_heading = doc.add_heading('Примечания:', level=2)
        notes_heading_run = notes_heading.runs[0]
        notes_heading_run.font.size = Pt(12)
        
        notes_para = doc.add_paragraph(notes)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    signature = doc.add_paragraph()
    signature.add_run('Дефектоскопист: ').bold = True
    signature.add_run('_' * 30)
    signature.add_run(f'  {inspector_name}')
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    docx_base64 = base64.b64encode(buffer.read()).decode('utf-8')
    
    return {
        'statusCode': 200,
        'headers': {
            'Content-Type': 'application/json',
            'Access-Control-Allow-Origin': '*'
        },
        'isBase64Encoded': False,
        'body': json.dumps({
            'success': True,
            'filename': f'conclusion_{joint_number}_{date}.docx',
            'document': docx_base64
        })
    }
